import os
from docx import Document
from docx.shared import Pt
def save_section_to_docx(section_title, text, image_path=None, file_path="data/output/thesis.docx"):
    if os.path.exists(file_path):
        doc = Document(file_path)
    else:
        doc = Document()
    doc.add_heading(section_title, level=1)
    paragraph = doc.add_paragraph(text)
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
    if image_path and os.path.exists(image_path):
        doc.add_picture(image_path)
    doc.save(file_path)
    return file_path